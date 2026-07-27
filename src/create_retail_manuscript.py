from __future__ import annotations

import os
from pathlib import Path
from PIL import Image, ImageOps, ImageDraw, ImageFont, ImageChops
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

BASE = Path('/mnt/data')
OUTDIR = BASE / 'sainsburys_retail_ai_outputs'
WORK = BASE / 'retail_manuscript_work'
WORK.mkdir(exist_ok=True)
FINAL = BASE / 'Retail_Decision_Intelligence_Full_Manuscript.docx'

# ---------- Figure assembly ----------
def trim(im: Image.Image, border=8) -> Image.Image:
    im = im.convert('RGB')
    bg = Image.new('RGB', im.size, 'white')
    diff = ImageChops.difference(im, bg)
    bbox = diff.getbbox()
    if bbox:
        im = im.crop((max(0,bbox[0]-border), max(0,bbox[1]-border), min(im.width,bbox[2]+border), min(im.height,bbox[3]+border)))
    return im

try:
    font_bold = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 44)
except Exception:
    font_bold = ImageFont.load_default()


def panel(files, outname, cols=2, target_w=1300, label_positions=None):
    images=[]
    for fn in files:
        im=trim(Image.open(OUTDIR/fn))
        scale=target_w/im.width
        im=im.resize((target_w, int(im.height*scale)), Image.Resampling.LANCZOS)
        images.append(im)
    rows=(len(images)+cols-1)//cols
    row_heights=[]
    for r in range(rows):
        row_heights.append(max(images[i].height for i in range(r*cols,min((r+1)*cols,len(images)))))
    gap=50
    margin=45
    canvas_w=cols*target_w+(cols-1)*gap+2*margin
    canvas_h=sum(row_heights)+(rows-1)*gap+2*margin
    canvas=Image.new('RGB',(canvas_w,canvas_h),'white')
    d=ImageDraw.Draw(canvas)
    labels='ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    y=margin
    for r in range(rows):
        x=margin
        for c in range(cols):
            i=r*cols+c
            if i>=len(images): break
            im=images[i]
            oy=y+(row_heights[r]-im.height)//2
            canvas.paste(im,(x,oy))
            # panel label with white backing
            d.rounded_rectangle((x+8,oy+8,x+74,oy+66),radius=8,fill='white',outline=(180,180,180),width=2)
            d.text((x+22,oy+8),labels[i],font=font_bold,fill=(10,35,80))
            x += target_w+gap
        y += row_heights[r]+gap
    path=WORK/outname
    canvas.save(path, quality=95)
    return path

fig2=panel(['01_daily_demand.png','02_stockout_by_category.png','03_promotion_multiplier.png'], 'Figure_2_simulation_benchmark.png', cols=2, target_w=1250)
fig3=panel(['04_forecast_wape.png','05_probabilistic_forecast.png','06_permutation_importance.png','07_segment_wape.png'], 'Figure_3_forecasting.png', cols=2, target_w=1250)
fig4=panel(['08_stockout_pr_curve.png','09_stockout_calibration.png'], 'Figure_4_stockout.png', cols=2, target_w=1250)
fig5=panel(['10_substitution_acceptance.png','11_substitution_tier_matrix.png'], 'Figure_5_substitution.png', cols=2, target_w=1250)
fig6=OUTDIR/'12_campaign_did.png'
fig7=OUTDIR/'13_inventory_outcomes.png'
fig1=BASE/'retail_decision_intelligence_infographic_design.png'

# ---------- Word helpers ----------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None:
            node=OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '1F4E79'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    rFonts=OxmlElement('w:rFonts'); rFonts.set(qn('w:ascii'),'Cambria'); rFonts.set(qn('w:hAnsi'),'Cambria'); rPr.append(rFonts)
    sz=OxmlElement('w:sz'); sz.set(qn('w:val'),'18'); rPr.append(sz)
    new_run.append(rPr)
    t=OxmlElement('w:t'); t.text=text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def set_columns(section, num=1, space_twips=360):
    sectPr=section._sectPr
    cols=sectPr.find(qn('w:cols'))
    if cols is None:
        cols=OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(space_twips))


def set_section_geometry(section):
    section.top_margin=Inches(0.62)
    section.bottom_margin=Inches(0.62)
    section.left_margin=Inches(0.65)
    section.right_margin=Inches(0.65)
    section.header_distance=Inches(0.25)
    section.footer_distance=Inches(0.25)


def set_run_font(run, name='Cambria', size=9.5, bold=None, italic=None, color=None):
    run.font.name=name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    if color: run.font.color.rgb=RGBColor(*color)


def add_body(text, first_line=True, italic=False, bold_lead=None):
    p=doc.add_paragraph(style='Body Text')
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_together=False
    if first_line:
        p.paragraph_format.first_line_indent=Inches(0.17)
    else:
        p.paragraph_format.first_line_indent=Inches(0)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); set_run_font(r,bold=True)
        r2=p.add_run(text[len(bold_lead):]); set_run_font(r2,italic=italic)
    else:
        r=p.add_run(text); set_run_font(r,italic=italic)
    return p


def add_heading(text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next=True
    r=p.add_run(text)
    return p


def add_equation(text):
    p=doc.add_paragraph(style='Equation')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); set_run_font(r, size=9.5, italic=True)
    return p


def add_figure(image_path, caption, width=7.0):
    # one-column continuous section for full-width figure
    sec=doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_geometry(sec); set_columns(sec,1)
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next=True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap=doc.add_paragraph(style='Caption')
    cap.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.keep_with_next=False
    r=cap.add_run(caption); set_run_font(r,size=8.2)
    sec2=doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_geometry(sec2); set_columns(sec2,2,340)
    return cap


def add_reference(authors_year_text, doi):
    p=doc.add_paragraph(style='References')
    r=p.add_run(authors_year_text + ' '); set_run_font(r,size=8.8)
    add_hyperlink(p, f'https://doi.org/{doi}', f'https://doi.org/{doi}')
    return p

# ---------- Document ----------
doc=Document()
sec0=doc.sections[0]
set_section_geometry(sec0); set_columns(sec0,1)

# Styles
styles=doc.styles
normal=styles['Normal']
normal.font.name='Cambria'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Cambria'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Cambria'); normal.font.size=Pt(9.5)
normal.paragraph_format.space_after=Pt(2.5); normal.paragraph_format.line_spacing=1.02
body=styles['Body Text']; body.font.name='Cambria'; body._element.rPr.rFonts.set(qn('w:ascii'),'Cambria'); body._element.rPr.rFonts.set(qn('w:hAnsi'),'Cambria'); body.font.size=Pt(9.5); body.paragraph_format.space_after=Pt(3); body.paragraph_format.line_spacing=1.03
for lvl,size in [(1,11.5),(2,10.5),(3,9.8)]:
    st=styles[f'Heading {lvl}']; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(20,45,80); st.paragraph_format.space_before=Pt(8 if lvl==1 else 5); st.paragraph_format.space_after=Pt(3)
styles['Caption'].font.name='Cambria'; styles['Caption']._element.rPr.rFonts.set(qn('w:ascii'),'Cambria'); styles['Caption']._element.rPr.rFonts.set(qn('w:hAnsi'),'Cambria'); styles['Caption'].font.size=Pt(8.2); styles['Caption'].font.italic=False; styles['Caption'].paragraph_format.space_after=Pt(5)
if 'Equation' not in styles:
    eq=styles.add_style('Equation', WD_STYLE_TYPE.PARAGRAPH)
else: eq=styles['Equation']
eq.font.name='Cambria Math'; eq.font.size=Pt(9.5); eq.paragraph_format.space_before=Pt(2); eq.paragraph_format.space_after=Pt(4)
if 'References' not in styles:
    rs=styles.add_style('References', WD_STYLE_TYPE.PARAGRAPH)
else: rs=styles['References']
rs.font.name='Cambria'; rs.font.size=Pt(8.8); rs.paragraph_format.left_indent=Inches(0.18); rs.paragraph_format.first_line_indent=Inches(-0.18); rs.paragraph_format.space_after=Pt(3); rs.paragraph_format.line_spacing=1.0

# Header/footer
for section in doc.sections:
    section.header.is_linked_to_previous=True
    section.footer.is_linked_to_previous=True
header=sec0.header.paragraphs[0]
header.alignment=WD_ALIGN_PARAGRAPH.CENTER
rh=header.add_run('RETAIL DECISION INTELLIGENCE | PETALCORIN')
set_run_font(rh,name='Arial',size=7.5,bold=True,color=(90,90,90))
footer=sec0.footer.paragraphs[0]
footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
rf=footer.add_run('Page '); set_run_font(rf,name='Arial',size=7.5,color=(100,100,100)); add_page_number(footer)

# Title
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(7)
r=p.add_run('Retail Decision Intelligence: An Evidence-Benchmarked Synthetic Framework Integrating Hierarchical Demand Forecasting, Latent-Demand Recovery, Stockout-Risk Prediction, Product-Substitution Ranking, Promotion Evaluation, and Probabilistic Inventory Optimisation')
set_run_font(r,name='Arial',size=19,bold=True,color=(12,35,72))

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Mark Ihrwell R. Petalcorin, PhD'); set_run_font(r,name='Arial',size=11,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('aAidea Ltd / Independent Applied AI Research, London, United Kingdom'); set_run_font(r,size=9.5,italic=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Correspondence: m.petalcorin@gmail.com'); set_run_font(r,size=9)

# Abstract
add_heading('Abstract',1)
abstract=(
'Grocery retail is a coupled stochastic system in which customer demand, pricing, promotions, inventory availability, substitution behaviour, and replenishment decisions interact across products, stores, and time. This study developed an evidence-benchmarked synthetic framework that connects predictive modelling with operational decisions rather than evaluating algorithms in isolation. A daily panel of 38,880 product-store observations was simulated for 18 products, four stores, six categories, and 540 days. Published estimates were used to calibrate promotional prevalence and discounting, promotional sales multipliers, a January plant-based campaign, overdispersed and zero-inflated sales, and supermarket stockout frequency. Latent demand was generated before inventory constraints, and observed sales were defined as the inventory-censored minimum of demand and available stock. Leakage-safe lag and rolling features were used to compare a seasonal-naive baseline with global gradient-boosted models trained on observed sales or a transparent stockout-corrected demand proxy. Quantile models supplied uncertainty estimates, a probabilistic classifier ranked stockout risk, a candidate-level model ranked product substitutions, a fixed-effects difference-in-differences model estimated campaign uplift, and forecast quantiles were translated into inventory policies. The demand-corrected global model reduced weighted absolute percentage error from 0.737 to 0.502, a relative reduction of 31.8%, with a bootstrap 95% confidence interval for mean absolute-error improvement of 3.232 to 3.784 units. The nominal 90% prediction interval achieved 93.8% empirical coverage. Stockout discrimination reached a receiver operating characteristic area under the curve of 0.747, although calibration remained imperfect. Attribute-aware substitution increased top-ranked acceptance from 27.0% under closest-price matching to 40.4%, while mean retained revenue increased from 1.34 to 1.93 simulated currency units per event. The campaign model estimated a 47.8% uplift against a programmed 57% effect, but uncertainty was substantial (robust p = 0.082). A 90th-quantile inventory policy increased service level from 66.7% to 95.8% and reduced total decision cost by 6.9%, while increasing waste. The results show, at the micro level, how stock censoring and product attributes alter inference, and, at the systems level, why forecasting, substitution, and inventory policies must be optimised jointly. Because all data are synthetic, the findings demonstrate methodological feasibility and trade-offs rather than retailer-specific causal effects.'
)
add_body(abstract, first_line=False)
p=doc.add_paragraph(); r=p.add_run('Keywords: '); set_run_font(r,bold=True,size=9); r2=p.add_run('retail analytics; demand forecasting; stockouts; substitution; promotion evaluation; probabilistic forecasting; inventory optimisation; explainable artificial intelligence'); set_run_font(r2,size=9)

# Main text to two columns
sec=doc.add_section(WD_SECTION.CONTINUOUS); set_section_geometry(sec); set_columns(sec,2,340)

add_heading('1. Introduction',1)
for text in [
'Grocery retail forecasting is not merely a time-series exercise. The quantity recorded at the checkout is the output of a constrained system: customers arrive with latent preferences, encounter prices and promotions, observe an assortment that may be incomplete, substitute or abandon when products are unavailable, and generate sales that feed subsequent replenishment decisions. The same observed sales series therefore reflects demand formation, merchandising, inventory control, and customer response. A model that predicts sales accurately but ignores these mechanisms can reinforce operational bias. In particular, a product that repeatedly stocks out may appear to have lower demand than it truly has because observed sales are capped by available inventory (Anupindi et al., 1998; Lee et al., 2016).',
'At the micro level, each product-store-day constitutes a stochastic decision unit. Demand is count-valued, frequently overdispersed, sometimes zero, and sensitive to temporal context, product attributes, store format, price, and promotion. Retail data also contain asymmetric errors: an under-forecast may produce lost margin, reduced availability, and customer dissatisfaction, whereas an over-forecast may create holding cost, markdowns, or food waste. The correct prediction target and loss function therefore depend on the downstream action. Probabilistic forecasts are especially valuable because the upper quantiles of demand can be translated directly into service-level and stocking decisions (Gneiting & Raftery, 2007; Spiliotis et al., 2021).',
'At the systems level, thousands of product-store series form a hierarchy linked through common seasonality, categories, locations, promotions, and supply constraints. The M5 competition demonstrated the scale and structure of this problem using 42,840 hierarchical retail series and showed the practical strength of global machine-learning approaches that pool information across related series (Makridakis et al., 2022a, 2022b; Ma & Fildes, 2022). Retail forecasting reviews likewise emphasise structural change, explanatory variables, omni-channel operations, and the need to evaluate whether forecast improvements translate into inventory and service improvements (Fildes et al., 2022a, 2022b).',
'Promotions and substitutions create additional feedback. Price promotions can be frequent and economically meaningful; a nationwide supermarket sample reported that 13.4% of products were promoted and promoted prices were 15.2% lower on average (Powell et al., 2016). Supermarket promotional displays and price signals can generate large sales multipliers, but effects vary by category, store context, and intervention design (Rosin et al., 2023; Trewern et al., 2022; Luick et al., 2024). When an item is unavailable, the retailer can either lose the sale or retain demand through a well-matched substitute. Consumer responses depend on product commitment, category differentiation, dominant attributes, and prior purchasing history (Fitzsimons, 2000; Hoang & Breugelmans, 2023).',
'This study developed a reproducible synthetic retail decision-intelligence framework with six connected analytical layers: latent-demand generation and censoring, global point forecasting, probabilistic forecasting, stockout-risk prediction, substitution ranking, promotion-effect estimation, and forecast-to-inventory optimisation. The objective was not to claim performance for a specific retailer, but to create a scientifically interpretable benchmark in which the true latent demand and programmed intervention effects were known. This permits direct examination of questions that are usually unobservable in real point-of-sale data: how much demand is hidden by stockouts, whether correcting the target reduces under-forecasting, whether ranked substitutes retain more demand, and whether predictive gains survive translation into operational costs.'
]: add_body(text)

add_figure(fig1, 'Figure 1. Integrated retail decision-intelligence framework. Product-store demand histories, inventory position, price and promotion signals, and product relationships are combined in a central analytical layer. The system produces probabilistic demand forecasts, stockout alerts, ranked substitutions, and inventory actions. At the micro level, the framework models individual product-store-day and candidate-substitute decisions. At the systems level, it links these decisions through feedback among availability, customer response, retained demand, service level, waste, and replenishment.', width=6.3)

add_heading('2. Materials and Methods',1)
add_heading('2.1 Study design and reproducibility',2)
for text in [
'A fully synthetic observational panel was generated with a fixed random seed of 42. The panel comprised four stores, 18 products, six product categories, and daily observations from 1 January 2024 through 23 June 2025. The resulting 72 product-store series contained 38,880 rows. The final 56 days, beginning 29 April 2025, were reserved as a chronological test set. All lagged and rolling features were shifted by at least one day, preventing future information from entering the predictors. This design follows the principle that forecasting models should be evaluated out of sample and in temporal order rather than through random splitting (Tashman, 2000).',
'The workflow was implemented in Python 3.13 using NumPy 2.3.5, pandas 2.2.3, scikit-learn 1.8.0, LightGBM 4.6.0, SciPy 1.17.0, statsmodels 0.14.6, and Matplotlib 3.10.8. The executed notebook, synthetic data, figures, and result files were generated from one deterministic pipeline.'
]: add_body(text)

add_heading('2.2 Evidence-benchmarked retail simulation',2)
add_body('Each product was assigned a category, brand, brand tier, base price, unit-cost ratio, baseline demand, perishability indicator, health score, package-size index, and price-elasticity parameter. Stores were assigned format, affluence, and footfall factors. Product-store baseline intensity was modified by weekday seasonality, weekend effects, annual seasonality, a small positive trend, holiday periods, temperature for beverages, promotions, and a January plant-based campaign. The expected demand intensity for product i in store s on day t was represented conceptually as a multiplicative system:')
add_equation('mu(i,s,t) = base(i,s) x weekly(t) x annual(t) x trend(t) x holiday(i,t) x promotion(i,t) x campaign(i,t) x weather(i,t).')
add_body('Demand was then sampled from a zero-inflated negative-binomial process. This choice reflects the overdispersion, intermittency, and excess zeros observed in retail sales data and highlighted in distributional analyses of M5 data (Ziel, 2022). The negative-binomial dispersion parameter was fixed at 5.0, while the zero probability decreased with the product-store baseline demand. Thus, slow-moving products generated more zero-demand days, whereas high-volume products remained more continuous.')
add_body('Promotion probability was set to 0.134, and discounts were sampled around 0.152 with bounds of 0.05 and 0.35, matching published supermarket prevalence and average price reduction (Powell et al., 2016). Promotional demand multipliers were sampled from 1.65 to 2.35, consistent with large sales responses reported for supermarket promotion mechanisms (Rosin et al., 2023). Plant-based products received an additional multiplier of 1.57 during January, benchmarked to the reported Veganuary-related increase in weekly plant-based sales (Trewern et al., 2022). Store affluence and format were retained to permit heterogeneous performance analysis, as real supermarket interventions can differ across store and socioeconomic contexts (Trewern et al., 2022; Luick et al., 2024).')

add_heading('2.3 Inventory censoring and latent-demand target',2)
add_body('Available stock was sampled around 1.85 times expected demand and was occasionally reduced by a disruption process whose probability increased for perishable and promoted products. Latent demand D(i,s,t) was generated before the stock constraint, available stock was A(i,s,t), and observed units sold Y(i,s,t) were defined as:')
add_equation('Y(i,s,t) = min[D(i,s,t), A(i,s,t)].')
add_body('A stockout occurred when D(i,s,t) exceeded A(i,s,t), and lost sales were D(i,s,t) - Y(i,s,t). This explicit ordering is central to the study. Once a stockout occurs, sales no longer equal demand; focal-product sales are right-censored, and substitute-product sales may be inflated (Anupindi et al., 1998). Historical supermarket studies report material stockout prevalence, supporting simulation calibration to an approximately 8%-10% operating range (Vasconcellos & Sampaio, 2009).')
add_body('A transparent demand proxy was constructed for stockout days as the maximum of observed sales, 1.15 times the shifted 28-day rolling mean, and available stock plus one-half of the shifted 28-day rolling standard deviation. Non-stockout days retained observed sales. This proxy was deliberately simple and auditable; it was used to test whether correcting the training target could reduce systematic under-forecasting.')

add_figure(fig2, 'Figure 2. Evidence-benchmarked simulation behaviour. (A) Aggregate daily latent demand and observed sales; the vertical gap represents demand censored by inventory. (B) Stockout rates by category, showing heterogeneity induced by product volume, perishability, promotion, and capacity disruption. (C) Observed promotion-to-non-promotion sales multipliers by category. The dashed reference indicates a two-fold multiplier. The synthetic panel contained a 13.8% promotion prevalence, 15.2% mean promotional discount, 9.2% stockout rate, and mean latent and observed demand of 16.06 and 15.40 units, respectively.')

add_heading('2.4 Feature engineering and demand forecasting',2)
add_body('Predictors included regular and realised price, discount, promotion, campaign, holiday, temperature, store affluence, perishability, health score, size index, calendar variables, encoded product and store identities, lags at 1, 7, 14, and 28 days, and shifted rolling means and standard deviations over 7 and 28 days. Three point-forecasting policies were compared: a seasonal-naive forecast equal to the previous week\'s sales, a global gradient-boosted model trained on observed sales, and an otherwise identical global model trained on the stockout-corrected proxy. Global models pooled all product-store series, allowing common responses to promotion, calendar, category, and recent demand to be learned across the hierarchy. This structure is consistent with the global bottom-up paradigm that performed strongly in the M5 retail competition (Makridakis et al., 2022b; Ma & Fildes, 2022).')
add_body('Models used 120 boosted trees, a learning rate of 0.05, 24 leaves, and row and feature subsampling of 0.85. Point predictions were clipped at zero. The true simulated latent demand was used only for evaluation. Metrics were mean absolute error (MAE), root mean squared error (RMSE), weighted absolute percentage error (WAPE), mean signed bias, and root mean squared logarithmic error. WAPE was defined as:')
add_equation('WAPE = sum |D - D_hat| / sum D.')
add_body('The difference in absolute error between seasonal-naive and corrected global predictions was bootstrapped to obtain a 95% confidence interval. Permutation importance was calculated by independently shuffling each feature in the held-out set and measuring the increase in MAE, an approach related to model-agnostic importance measures developed for tree ensembles (Breiman, 2001). Performance was also stratified by store format.')

add_heading('2.5 Probabilistic forecasting',2)
add_body('Separate gradient-boosted quantile models estimated the 0.05, 0.50, 0.90, and 0.95 conditional demand quantiles. The 0.05 and 0.95 predictions defined a nominal 90% interval. Evaluation included empirical coverage and mean interval width. Probabilistic forecasts were treated as decision distributions rather than decorative uncertainty bands: calibration determines whether a quantile corresponds to the intended service probability, and proper evaluation is required to align probabilistic predictions with decisions (Gneiting & Raftery, 2007; Makridakis et al., 2022c; Chen et al., 2022).')

add_heading('2.6 Stockout-risk prediction',2)
add_body('A gradient-boosted binary classifier predicted whether latent demand would exceed known available stock. Inputs combined demand-history features, product and store attributes, prices and promotions, and available stock. Discrimination was evaluated by receiver operating characteristic area under the curve (ROC-AUC) and average precision. Because stockouts were uncommon relative to non-stockouts, the precision-recall curve was emphasised; precision-recall analysis directly reflects positive-prediction reliability under class imbalance (Saito & Rehmsmeier, 2015). Probability quality was assessed by the Brier score and a ten-bin calibration curve (Brier, 1950). An operational threshold was chosen by maximising F2, which weights recall more strongly than precision and therefore reflects a setting where missing a genuine stockout is more costly than investigating a false alert.')

add_heading('2.7 Product-substitution ranking',2)
add_body('A separate simulation created 2,500 out-of-stock events. Each event presented four candidate substitutes and an option to reject all candidates. Candidate-level features were price gap, package-size gap, health-score gap, same brand, same brand tier, same category, customer affinity, prior purchase, candidate price, and display position. Utilities were converted to a five-way softmax choice distribution. The absolute event acceptance level was a transparent scenario assumption rather than a literature-fixed constant. The mechanism, however, was evidence-based: dominant-attribute matching and prior purchase can materially increase online grocery substitution acceptance (Hoang & Breugelmans, 2023).')
add_body('Events 0-1,999 were used for training and events 2,000-2,499 for testing, ensuring that all four candidates from one event remained in the same split. A class-balanced logistic model estimated candidate acceptance probabilities. One candidate per event was selected by the highest predicted probability and compared with closest-price matching and a purchase-history heuristic. Primary outcomes were top-1 acceptance and retained revenue, defined as candidate price multiplied by the indicator that the recommended candidate was accepted.')

add_heading('2.8 Promotion-effect estimation',2)
add_body('Weekly sales from December 2024 through February 2025 were aggregated by store and category. Plant-based products formed the treated group; Dairy and Household categories were controls. A log-linear difference-in-differences model included treated, post-January, their interaction, store fixed effects, and week fixed effects:')
add_equation('log(1 + sales) = beta0 + beta1 treated + beta2 post + beta3(treated x post) + store effects + week effects + error.')
add_body('The campaign uplift was exp(beta3) - 1. Heteroscedasticity-consistent HC3 standard errors were reported. Difference-in-differences can reduce confounding by common shocks and persistent group differences, but serial correlation, pre-trend violations, intervention heterogeneity, and spillovers require careful treatment in real applications (Bertrand et al., 2004). Because the simulation contained a programmed 57% January effect, the estimate could be compared with a known benchmark.')

add_heading('2.9 Forecast-to-inventory optimisation',2)
add_body('Two replenishment policies were evaluated on the held-out period. The seasonal-naive policy stocked the previous week\'s sales plus a square-root safety buffer. The probabilistic policy stocked the ceiling of the predicted 90th demand quantile. For stock quantity Q and realised demand D, fulfilled units were min(D,Q), lost units were max(D-Q,0), and leftover units were max(Q-D,0). Perishable products incurred a 40% spoilage fraction of leftovers and non-perishable products 5%; remaining leftovers incurred a 3% holding charge. Total decision cost combined lost margin, waste cost, and holding cost. This is a data-driven newsvendor formulation: the operational objective is a quantile of the conditional demand distribution determined by underage and overage costs, not necessarily the mean forecast (Qin et al., 2011; Huber et al., 2019). Outcomes were service level, lost units, waste units, total decision cost, retained-margin proxy, and average stock.')

add_heading('3. Results',1)
add_heading('3.1 Synthetic panel reproduced the intended retail mechanisms',2)
add_body('The generated panel contained 38,880 observations and 72 product-store series. Promotions occurred on 13.8% of rows, with a mean conditional discount of 15.2%. The overall stockout rate was 9.2%, and 13.3% of rows had zero observed sales. Mean latent demand was 16.056 units, whereas mean observed sales were 15.401, producing a mean censoring gap of 0.656 units per product-store-day. Aggregated time series showed that observed sales tracked demand on most days but diverged during high-demand or low-capacity episodes (Figure 2A). Category stockout rates ranged from approximately 8% to above 10%, and promotions generated approximately two-fold observed sales across categories (Figure 2B-C).')
add_body('At the micro level, the simulation therefore created two distinct sources of variability: stochastic demand around a conditional mean and a one-sided measurement process that truncated demand at available stock. At the systems level, the truncation affected the data used by subsequent forecasting and replenishment modules, creating the possibility of a self-reinforcing loop in which stockouts suppress recorded sales, suppressed sales lower forecasts, and lower forecasts perpetuate insufficient stock.')

add_heading('3.2 Global models improved latent-demand forecasting',2)
add_body('The seasonal-naive baseline achieved an MAE of 10.950 units, RMSE of 16.024, and WAPE of 0.737. The global model trained on observed sales reduced MAE to 7.467 and WAPE to 0.502. Training on the corrected demand proxy produced an MAE of 7.462, RMSE of 10.921, and WAPE of 0.502. Relative to seasonal naive, the corrected model reduced WAPE by 31.8% and absolute error by 3.487 units on average. The 95% bootstrap confidence interval for the mean absolute-error improvement was 3.232 to 3.784 units.')
add_body('The corrected and observed-sales models had similar aggregate WAPE, but the corrected model reduced mean under-forecast bias from -1.276 to -0.909 units. This difference is operationally important even when aggregate accuracy changes little: bias accumulates across replenishment cycles, and persistent negative bias can increase stockout exposure. The probabilistic 90% interval achieved 93.8% empirical coverage with a mean width of 31.30 units, indicating conservative but useful uncertainty quantification.')
add_body('Permutation analysis identified the shifted 28-day rolling mean as the dominant predictor, followed by discount, promotion, weekday, category, and recent variability. This hierarchy is mechanistically plausible: the rolling mean captures local product-store demand intensity, while promotion and discount alter short-term demand conditional on that baseline. WAPE differed by store format, from 0.473 for superstores to 0.535 for urban stores and 0.562 for convenience stores, showing that one global model can still have heterogeneous local performance.')

add_figure(fig3, 'Figure 3. Demand-forecasting results. (A) WAPE for the seasonal-naive, observed-sales global, and demand-corrected global models when evaluated against latent demand. (B) Example probabilistic forecast for one product-store series, showing median prediction and nominal 90% interval. (C) Held-out permutation importance expressed as increase in MAE after shuffling. (D) WAPE stratified by store format. The demand-corrected model reduced WAPE by 31.8% relative to seasonal naive; the nominal 90% interval covered 93.8% of observations.')

add_heading('3.3 Stockout risk was predictable but not fully calibrated',2)
add_body('The stockout classifier achieved ROC-AUC 0.747 and average precision 0.285 against a test prevalence of 0.084. The Brier score was 0.186. Maximisation of F2 selected a threshold of 0.451. At this threshold the model produced 229 true positives, 109 false negatives, 1,125 false positives, and 2,569 true negatives, corresponding to recall of 67.8% and precision of 16.9%. The threshold therefore recovered more than two-thirds of stockouts but generated a substantial review burden.')
add_body('The precision-recall curve showed that precision declined as recall approached one, illustrating the practical cost of broad alerting. Calibration was directionally ordered: observed stockout frequency increased from approximately 2%-3% in lower-risk bins to 27.7% in the highest-risk bin. However, the model systematically overpredicted absolute probabilities in most bins. Thus, it was more reliable as a ranking tool than as an uncalibrated probability engine. In deployment, post-hoc calibration or retraining against changing inventory processes would be necessary before probabilities were used directly in expected-cost calculations.')

add_figure(fig4, 'Figure 4. Stockout-risk evaluation. (A) Precision-recall curve with the F2-selected operating threshold and prevalence baseline. (B) Probability calibration by risk decile compared with perfect calibration. The model achieved ROC-AUC 0.747 and average precision 0.285. Ranking was informative, but predicted probabilities exceeded observed stockout rates in most bins.')

add_heading('3.4 Attribute-aware substitution retained more demand',2)
add_body('Across all simulated events, at least one candidate was accepted in 73.9% of cases. In the held-out event set, closest-price matching achieved 27.0% top-1 acceptance, and the purchase-history heuristic achieved 32.2%. The attribute-aware model achieved 40.4%, an absolute gain of 13.4 percentage points over closest-price matching and 8.2 points over purchase history alone. Mean retained revenue increased from 1.341 under closest-price matching to 1.511 under the history heuristic and 1.927 under the attribute-aware model.')
add_body('The accepted-flow matrix showed a strong preference for substitutes within the same brand tier: 88% of accepted alternatives for premium sources were premium, and 86% for standard sources were standard. Value products were more diffuse, but 60% of accepted substitutes remained in the value tier. At the micro level, candidate acceptance therefore depended on a bundle of similarity and preference signals rather than price distance alone. At the system level, improved substitution reduced the fraction of stockouts that converted directly into lost demand, which can change the optimal inventory mix and the apparent sales of substitute items.')

add_figure(fig5, 'Figure 5. Product-substitution results. (A) Top-ranked acceptance under closest-price, purchase-history, and attribute-aware machine-learning policies. (B) Conditional distribution of accepted candidate brand tiers by source tier. The attribute-aware policy increased top-1 acceptance from 27.0% to 40.4% and increased the retained-revenue proxy by 43.8% relative to closest-price matching.')

add_heading('3.5 The campaign model recovered the programmed direction with uncertainty',2)
add_body('Plant-based indexed sales increased sharply during January relative to the control categories, then declined after the intervention period (Figure 6). The difference-in-differences interaction coefficient was 0.391 with a robust standard error of 0.225, corresponding to an estimated uplift of 47.8%. This estimate recovered the direction and approximate magnitude of the programmed 57% campaign multiplier but did not reach the conventional 0.05 significance threshold (robust p = 0.082).')
add_body('The result illustrates why a large visual change is not equivalent to a precise causal estimate. The effective sample size was determined by the number of store-category-week cells and by within-panel dependence, not by the original number of daily product rows. Moreover, the simulation did not model all real promotional complications, including anticipatory purchasing, post-promotion dips, assortment changes, local competitor actions, or interference between categories. The estimate should therefore be viewed as a design demonstration rather than evidence for a real intervention.')

add_figure(fig6, 'Figure 6. Indexed weekly sales around the simulated January plant-based campaign. Values were normalised to the first pre-campaign week. The difference-in-differences model estimated a 47.8% incremental uplift compared with a programmed 57% multiplier; robust p = 0.082.')

add_heading('3.6 Probabilistic inventory policy improved service at a waste cost',2)
add_body('The seasonal-naive stock policy achieved a service level of 66.7%, lost 19,981 units, generated 8,024.75 waste units, and incurred a total decision cost of 63,084.13 simulated currency units. The 90th-quantile policy increased service level to 95.8% and reduced lost units to 2,515. Retained margin increased from 50,532.43 to 54,904.42, and total decision cost decreased to 58,712.15, a reduction of 6.9%.')
add_body('The improvement required more inventory. Average stock rose from 16.22 to 27.81 units, and waste increased to 17,079 units. The policy therefore did not dominate on every operational dimension. Instead, it moved the system along a service-waste frontier: high quantiles protected availability and margin but increased overage for perishable goods. This is the central systems-level result. Forecast accuracy alone cannot identify the appropriate policy; the chosen quantile must reflect product-specific underage costs, waste, shelf life, lead time, substitution potential, and service targets.')

add_figure(fig7, 'Figure 7. Normalised inventory outcomes under seasonal-naive and probabilistic policies. Lower values are preferable. The 90th-quantile policy sharply reduced lost units and reduced total decision cost by 6.9%, but increased waste because it stocked more aggressively. Service level increased by 29.2 percentage points, from 66.7% to 95.8%.')

add_heading('4. Discussion',1)
add_heading('4.1 Micro-level interpretation: the product-store-day as a constrained observation',2)
for text in [
'The framework shows that a retail observation has both a behavioural and an operational origin. Latent demand was generated from product, store, temporal, weather, price, and promotional determinants, but observed sales were filtered through inventory. Consequently, zero sales could mean zero demand or zero availability; a low sales value could mean low preference or insufficient stock. This distinction is statistically analogous to censoring and missing-not-at-random processes: the probability that demand is hidden is greatest exactly when demand is high relative to capacity. Naively treating observed sales as the target therefore creates a systematic error mechanism rather than symmetric noise.',
'The demand-proxy experiment produced only a small change in aggregate WAPE relative to the observed-sales global model, but it reduced negative bias. This is consistent with the idea that many non-stockout observations still identify the demand function, while a minority of censored observations disproportionately affect replenishment risk. More sophisticated recovery could estimate stockout time, customer arrivals, and substitution flows jointly. Anupindi et al. (1998) formalised the bias created by stockout substitution, and field evidence has shown that attribute-based models can improve both demand estimation and inventory planning under extensive stockouts (Lee et al., 2016). The present proxy should be interpreted as a transparent first step, not a preferred estimator.',
'The substitution results further illustrate the limits of univariate heuristics. Closest price ignores brand architecture, package equivalence, health characteristics, category differentiation, and customer history. Acceptance increased when these signals were combined, agreeing with experimental findings that dominant-attribute and past-purchase matching improve online grocery substitutions (Hoang & Breugelmans, 2023). The tier matrix suggests that asymmetric preferences matter: premium shoppers may reject a value replacement even when price is close, whereas a value shopper may accept movement across tiers when the remaining attributes fit.'
]: add_body(text)

add_heading('4.2 Systems-level interpretation: prediction, action, and feedback',2)
for text in [
'A retail decision stack is a feedback control system. Forecasts set inventory; inventory determines availability; availability determines observed sales and substitutions; those observations become the next training data; promotions change both demand and the value of availability; and service failures alter future customer behaviour. Optimising one module in isolation can therefore produce unintended consequences. An aggressive forecast can reduce lost sales but increase waste. A strong substitution system can buffer stockouts, changing the underage cost and optimal order quantity. A campaign can appear successful because it shifts demand from control categories or because promoted products were better stocked, not solely because preferences changed.',
'The inventory experiment made this coupling explicit. The probabilistic policy reduced decision cost and increased service, but it did so through substantially higher stock and waste. Newsvendor theory represents this trade-off through a critical quantile determined by underage and overage costs (Qin et al., 2011). Data-driven newsvendor research demonstrates that conditional quantile learning can translate rich features into improved decisions without assuming a fixed demand distribution (Huber et al., 2019). In practice, the appropriate quantile should vary by product: fresh food with short shelf life, high substitution acceptance, and low lost margin may require a lower quantile than a high-margin, non-perishable product with few substitutes.',
'The stockout classifier also illustrates the difference between ranking and action. ROC-AUC of 0.747 indicated useful discrimination, but the selected high-recall threshold produced low precision. Whether this is acceptable depends on the intervention. A low-cost automated replenishment review may tolerate many alerts; expensive expedited shipments may require a higher threshold. Calibration is essential when predicted probabilities are multiplied by costs. The observed overprediction means that a decision rule using uncalibrated risk would overstate expected stockout cost and could systematically overstock.',
'Promotion evaluation must be integrated with forecasting and inventory rather than treated as a separate marketing report. Promotions create nonstationarity, alter price elasticity, affect stockout probability, and may shift demand across products and time. Real supermarket natural experiments demonstrate that effects vary across categories and store populations (Trewern et al., 2022; Luick et al., 2024). A production system should therefore forecast the counterfactual baseline, estimate heterogeneous incremental effects, and ensure that promoted inventory can support the expected uplift.'
]: add_body(text)

add_heading('4.3 Relation to prior retail forecasting research',2)
add_body('The 31.8% WAPE reduction relative to a seasonal-naive baseline is qualitatively consistent with the strong performance of global tree-based methods in the M5 competition, but direct numerical comparison would be inappropriate because the present dataset, horizon, hierarchy, and metric scaling differ (Makridakis et al., 2022a, 2022b). The main transferable finding is structural: cross-series learning can exploit shared demand drivers and product-store identities, while probabilistic models provide the distributional information needed for safety-stock decisions. M5 uncertainty analyses likewise emphasise overdispersion, zero demand, and quantile evaluation (Makridakis et al., 2022c; Ziel, 2022; Spiliotis et al., 2021).')
add_body('The framework also supports the broader recommendation that forecast models should be judged by decision value. Fildes et al. (2022a, 2022b) noted that retail forecasting practice must account for organisational process, structural breaks, machine learning, and the uncertain link between forecast accuracy and stock outcomes. Here, the corrected model had almost identical WAPE to the observed-sales model but lower negative bias, and the probabilistic policy generated better service and cost despite higher waste. These outcomes would be invisible if evaluation stopped at one aggregate point-forecast metric.')

add_heading('4.4 Limitations',2)
for text in [
'First, synthetic data provide complete ground truth but cannot reproduce the full joint distribution of a real retailer. Customer heterogeneity, basket structure, store geography, supplier lead times, replenishment schedules, shelf capacity, online picking, competitor actions, inflation, and long-term preference change were simplified or absent. Evidence-benchmarked parameter ranges improve plausibility but do not establish external validity.',
'Second, the stockout-correction rule was heuristic. It assumed daily availability and did not observe the time within the day at which inventory reached zero. Real demand recovery should use intraday inventory events, replenishment receipts, shelf audits, substitution transactions, and models that account for informative censoring.',
'Third, the test used one 56-day holdout. Rolling-origin evaluation across seasons, disruptions, and promotional regimes would provide stronger evidence of stability (Tashman, 2000). Hyperparameters were not extensively tuned, and uncertainty intervals were evaluated mainly through coverage rather than a full suite of proper scoring rules.',
'Fourth, the substitution model represented one accepted candidate per event and did not simulate full baskets, product complements, customer lifetime effects, or fairness constraints. Absolute acceptance was an assumption. Fifth, the campaign analysis used a short window and a simple treated-control structure. The p-value of 0.082 indicates that the estimate was imprecise, and a real causal claim would require stronger pre-trend, interference, and robustness analyses. Sixth, the inventory cost function was illustrative; all financial quantities were proxies, and the waste fractions were scenario parameters rather than empirically estimated retailer values.'
]: add_body(text)

add_heading('4.5 Production and governance implications',2)
add_body('A production implementation should begin with a unified product-store-day feature layer that preserves event time, inventory state, price and promotion history, product metadata, and substitution outcomes. Baselines should be maintained alongside global point and quantile models. Rolling backtests should report not only mean accuracy but bias, calibration, category and store performance, high-demand-tail error, and operational cost. Stockout probabilities should be calibrated and monitored for drift. Substitution policies should be evaluated through controlled experiments with acceptance, retained revenue, customer satisfaction, and healthy-choice or supplier-fairness outcomes.')
add_body('Decision ownership is equally important. Forecasting teams, supply-chain operators, commercial teams, digital product managers, and store colleagues optimise different objectives. A model card should specify the target, known censoring, prediction horizon, intended decision, threshold or quantile, cost assumptions, fallback rules, and monitoring cadence. Human override data should be retained as a source of learning rather than discarded. At the systems level, successful retail AI is therefore socio-technical: predictive models, operational constraints, experiments, user interfaces, and governance must be designed as one decision process.')

add_heading('5. Conclusion',1)
add_body('This study developed an evidence-benchmarked synthetic retail decision-intelligence framework that connected latent-demand modelling, global forecasting, uncertainty, stockout risk, substitution, campaign evaluation, and inventory optimisation. The corrected global model reduced WAPE by 31.8% relative to seasonal naive and reduced under-forecast bias. Probabilistic intervals achieved 93.8% coverage. Stockout risk was rankable but required calibration, attribute-aware substitution increased top-1 acceptance by 13.4 percentage points, and a 90th-quantile stocking policy improved service and reduced total cost while increasing waste. The campaign analysis recovered the programmed direction but remained statistically uncertain.')
add_body('The principal scientific implication is that retail sales are constrained system outputs, not direct measurements of unconstrained demand. At the micro level, inventory censoring, price, promotion, and product similarity shape each observation and choice. At the systems level, forecasts alter stock, stock alters sales, substitution alters retained demand, and these feedbacks determine service, waste, and cost. Retail machine learning should therefore be evaluated as a decision system, with calibrated uncertainty and explicit operational trade-offs, rather than as a collection of isolated predictive models.')

add_heading('Data and Code Availability',1)
add_body('The analysis used only synthetic data. The executed Jupyter notebook, generated synthetic panel, figures, and result tables accompany this manuscript as project artifacts. No customer, employee, supplier, or retailer-confidential data were used.', first_line=False)
add_heading('Conflict of Interest',1)
add_body('The author declares no competing financial interest. This work is a methodological portfolio study and is not an analysis of, or commissioned by, any named retailer.', first_line=False)

add_heading('References',1)
refs=[
('Anupindi, R., Dada, M., & Gupta, S. (1998). Estimation of consumer demand with stock-out based substitution: An application to vending machine products. Marketing Science, 17(4), 406-423.','10.1287/mksc.17.4.406'),
('Bertrand, M., Duflo, E., & Mullainathan, S. (2004). How much should we trust differences-in-differences estimates? The Quarterly Journal of Economics, 119(1), 249-275.','10.1162/003355304772839588'),
('Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.','10.1023/A:1010933404324'),
('Brier, G. W. (1950). Verification of forecasts expressed in terms of probability. Monthly Weather Review, 78(1), 1-3.','10.1175/1520-0493(1950)078<0001:VOFEIT>2.0.CO;2'),
('Chen, Z., Gaba, A., Tsetlin, I., & Winkler, R. L. (2022). Evaluating quantile forecasts in the M5 uncertainty competition. International Journal of Forecasting, 38(4), 1531-1545.','10.1016/j.ijforecast.2022.03.004'),
('Fildes, R., Ma, S., & Kolassa, S. (2022a). Retail forecasting: Research and practice. International Journal of Forecasting, 38(4), 1283-1318.','10.1016/j.ijforecast.2019.06.004'),
('Fildes, R., Kolassa, S., & Ma, S. (2022b). Post-script-Retail forecasting: Research and practice. International Journal of Forecasting, 38(4), 1319-1324.','10.1016/j.ijforecast.2021.09.012'),
('Fitzsimons, G. J. (2000). Consumer response to stockouts. Journal of Consumer Research, 27(2), 249-266.','10.1086/314323'),
('Gneiting, T., & Raftery, A. E. (2007). Strictly proper scoring rules, prediction, and estimation. Journal of the American Statistical Association, 102(477), 359-378.','10.1198/016214506000001437'),
('Hoang, D., & Breugelmans, E. (2023). "Sorry, the product you ordered is out of stock": Effects of substitution policy in online grocery retailing. Journal of Retailing, 99(1), 26-45.','10.1016/j.jretai.2022.06.006'),
('Huber, J., Müller, S., Fleischmann, M., & Stuckenschmidt, H. (2019). A data-driven newsvendor problem: From data to decision. European Journal of Operational Research, 278(3), 904-915.','10.1016/j.ejor.2019.04.043'),
('Lee, J., Gaur, V., Muthulingam, S., & Swisher, G. F. (2016). Stockout-based substitution and inventory planning in textbook retailing. Manufacturing & Service Operations Management, 18(1), 104-121.','10.1287/msom.2015.0551'),
('Luick, M., Bandy, L., Piernas, C., Jebb, S. A., & Pechey, R. (2024). Do promotions of healthier or more sustainable foods increase sales? Findings from three natural experiments in UK supermarkets. BMC Public Health, 24, 1658.','10.1186/s12889-024-19080-x'),
('Ma, S., & Fildes, R. (2022). The performance of the global bottom-up approach in the M5 accuracy competition: A robustness check. International Journal of Forecasting, 38(4), 1492-1499.','10.1016/j.ijforecast.2021.09.002'),
('Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022a). The M5 competition: Background, organization, and implementation. International Journal of Forecasting, 38(4), 1325-1336.','10.1016/j.ijforecast.2021.07.007'),
('Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022b). M5 accuracy competition: Results, findings, and conclusions. International Journal of Forecasting, 38(4), 1346-1364.','10.1016/j.ijforecast.2021.11.013'),
('Makridakis, S., Spiliotis, E., Assimakopoulos, V., Chen, Z., Gaba, A., Tsetlin, I., & Winkler, R. L. (2022c). The M5 uncertainty competition: Results, findings and conclusions. International Journal of Forecasting, 38(4), 1365-1385.','10.1016/j.ijforecast.2021.10.009'),
('Powell, L. M., Kumanyika, S. K., Isgor, Z., Rimkus, L., Zenk, S. N., & Chaloupka, F. J. (2016). Price promotions for food and beverage products in a nationwide sample of food stores. Preventive Medicine, 86, 106-113.','10.1016/j.ypmed.2016.01.011'),
('Qin, Y., Wang, R., Vakharia, A. J., Chen, Y., & Seref, M. M. H. (2011). The newsvendor problem: Review and directions for future research. European Journal of Operational Research, 213(2), 361-374.','10.1016/j.ejor.2010.11.024'),
('Rosin, M., Young, L., Jiang, Y., Vandevijvere, S., Waterlander, W., Mackay, S., & Ni Mhurchu, C. (2023). Product promotional strategies in supermarkets and their effects on sales: A case study of breakfast cereals and drinks in New Zealand. Nutrition & Dietetics, 80(5), 463-471.','10.1111/1747-0080.12800'),
('Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLOS ONE, 10(3), e0118432.','10.1371/journal.pone.0118432'),
('Spiliotis, E., Makridakis, S., Kaltsounis, A., & Assimakopoulos, V. (2021). Product sales probabilistic forecasting: An empirical evaluation using the M5 competition data. International Journal of Production Economics, 240, 108237.','10.1016/j.ijpe.2021.108237'),
('Tashman, L. J. (2000). Out-of-sample tests of forecasting accuracy: An analysis and review. International Journal of Forecasting, 16(4), 437-450.','10.1016/S0169-2070(00)00065-0'),
('Trewern, J., Chenoweth, J., Christie, I., & Halevy, S. (2022). Does promoting plant-based products in Veganuary lead to increased sales, and a reduction in meat sales? A natural experiment in a supermarket setting. Public Health Nutrition, 25(11), 3204-3214.','10.1017/S1368980022001914'),
('Vasconcellos, L. H. R., & Sampaio, M. (2009). The stockouts study: An examination of the extent and the causes in the São Paulo supermarket sector. Brazilian Administration Review, 6(3), 263-279.','10.1590/S1807-76922009000300007'),
('Ziel, F. (2022). M5 competition uncertainty: Overdispersion, distributional forecasting, GAMLSS, and beyond. International Journal of Forecasting, 38(4), 1546-1554.','10.1016/j.ijforecast.2021.09.008'),
]
for txt,doi in refs: add_reference(txt,doi)

# Supplementary material - single column
secS=doc.add_section(WD_SECTION.CONTINUOUS); set_section_geometry(secS); set_columns(secS,1)
add_heading('Supplementary Material',1)
add_body('All tables are placed in the Supplementary Material so that the main manuscript remains figure-led and compatible with a two-column journal layout. Values are derived from the executed synthetic notebook.', first_line=False)

# Table helper
def add_table_title(title):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(3); p.paragraph_format.keep_with_next=True
    r=p.add_run(title); set_run_font(r,name='Arial',size=9.5,bold=True,color=(20,45,80))

def add_table(headers, rows, widths=None, font_size=8.2):
    table=doc.add_table(rows=1, cols=len(headers))
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style='Table Grid'
    table.autofit=True
    hdr=table.rows[0].cells
    for j,h in enumerate(headers):
        hdr[j].text=str(h); set_cell_shading(hdr[j],'D9EAF7'); hdr[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(hdr[j])
        for run in hdr[j].paragraphs[0].runs: set_run_font(run,name='Arial',size=font_size,bold=True)
    for row in rows:
        cells=table.add_row().cells
        for j,val in enumerate(row):
            cells[j].text=str(val); cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cells[j])
            for p in cells[j].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for run in p.runs: set_run_font(run,size=font_size)
    return table

add_table_title('Supplementary Table S1. Literature-benchmarked simulation parameters.')
rows=[
('Promotion prevalence','0.134','Powell et al. (2016)','Mean supermarket promotion prevalence'),
('Mean promotional discount','0.152, truncated 0.05-0.35','Powell et al. (2016)','Promoted prices approximately 15.2% lower'),
('Promotion demand multiplier','Uniform 1.65-2.35','Rosin et al. (2023)','Promotions produced approximately 2-2.5-fold sales'),
('January plant-based multiplier','1.57','Trewern et al. (2022)','Reported 57% increase during intervention'),
('Target stockout range','Approximately 8%-10%','Vasconcellos & Sampaio (2009)','Historical supermarket stockout estimates'),
('Demand distribution','Zero-inflated negative binomial','Ziel (2022)','Overdispersion, intermittency, and zero demand'),
('Substitution mechanism','Attribute and purchase-history utility','Hoang & Breugelmans (2023)','Dominant attribute and past purchase affect acceptance'),
('Absolute substitution acceptance','Scenario-generated','Not fixed to one study','No universal acceptance probability assumed'),
]
add_table(['Parameter','Simulation value','Peer-reviewed source','Rationale'],rows,font_size=7.8)

add_table_title('Supplementary Table S2. Synthetic dataset characteristics.')
rows=[
('Random seed','42'),('Date range','1 January 2024-23 June 2025'),('Held-out test period','29 April 2025-23 June 2025'),('Stores','4'),('Products','18'),('Categories','6'),('Product-store series','72'),('Rows','38,880'),('Promotion prevalence','13.8%'),('Mean promotional discount','15.2%'),('Stockout rate','9.2%'),('Zero-sales rate','13.3%'),('Mean latent demand','16.056 units'),('Mean observed sales','15.401 units'),('Mean censoring gap','0.656 units')]
add_table(['Characteristic','Value'],rows,font_size=8.2)

doc.add_page_break()
add_table_title('Supplementary Table S3. Point-forecast performance on future latent demand.')
fm=pd.read_csv(OUTDIR/'forecast_metrics.csv')
rows=[]
for _,r in fm.iterrows(): rows.append((r['Model'],f"{r['MAE']:.3f}",f"{r['RMSE']:.3f}",f"{r['WAPE']:.3f}",f"{r['Bias']:.3f}",f"{r['RMSLE']:.3f}"))
add_table(['Model','MAE','RMSE','WAPE','Bias','RMSLE'],rows,font_size=8.0)

add_table_title('Supplementary Table S4. Stockout, substitution, promotion, and inventory outcomes.')
rows=[
('Stockout classifier','ROC-AUC','0.747'),('Stockout classifier','Average precision','0.285'),('Stockout classifier','Brier score','0.186'),('Stockout classifier','Prevalence','0.084'),('Stockout threshold','F2-selected probability','0.451'),('Stockout threshold','Precision','0.169'),('Stockout threshold','Recall','0.678'),
('Substitution: closest price','Top-1 acceptance','27.0%'),('Substitution: purchase history','Top-1 acceptance','32.2%'),('Substitution: attribute-aware ML','Top-1 acceptance','40.4%'),('Substitution: attribute-aware ML','Mean retained revenue','1.927'),
('Promotion model','Estimated uplift','47.8%'),('Promotion model','Robust p-value','0.082'),
('Seasonal-naive inventory','Service level','66.7%'),('Probabilistic inventory','Service level','95.8%'),('Probabilistic inventory','Decision-cost reduction','6.9%'),('Probabilistic inventory','Service-level increase','29.2 percentage points')]
add_table(['Analysis','Metric','Value'],rows,font_size=8.0)

add_table_title('Supplementary Table S5. Inventory-policy comparison.')
inv=pd.read_csv(OUTDIR/'inventory_policy_results.csv')
rows=[]
for _,r in inv.iterrows():
    rows.append((r['Policy'],f"{r['Service level']:.3f}",f"{r['Lost units']:,.0f}",f"{r['Waste units']:,.2f}",f"{r['Total decision cost']:,.2f}",f"{r['Retained margin proxy']:,.2f}",f"{r['Average stock']:.2f}"))
add_table(['Policy','Service level','Lost units','Waste units','Decision cost','Retained margin','Average stock'],rows,font_size=7.7)

add_table_title('Supplementary Table S6. Forecast performance by store format.')
seg=pd.read_csv(OUTDIR/'segment_performance.csv')
rows=[]
for _,r in seg.iterrows(): rows.append((r['store_format'],f"{int(r['n']):,}",f"{r['WAPE']:.3f}",f"{r['Bias']:.3f}"))
add_table(['Store format','Test observations','WAPE','Bias'],rows,font_size=8.2)

# Add alt text to images by filename via a11y tool later; set core properties
doc.core_properties.title='Retail Decision Intelligence: Evidence-Benchmarked Synthetic Framework'
doc.core_properties.subject='Retail demand forecasting, stockouts, substitution, promotion evaluation, and inventory optimisation'
doc.core_properties.author='Mark Ihrwell R. Petalcorin'
doc.core_properties.keywords='retail analytics, forecasting, stockout, substitution, inventory optimisation, machine learning'
doc.core_properties.comments='Generated from an executed synthetic retail analytics notebook.'

# Ensure section geometry and header/footer linkage for all sections
for s in doc.sections:
    set_section_geometry(s)
    s.header.is_linked_to_previous=True
    s.footer.is_linked_to_previous=True

# save
doc.save(FINAL)
print(FINAL)
